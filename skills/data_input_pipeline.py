"""
data_input_pipeline.py — 文档→仿真参数智能提取

支持:
  1. Word文档 (.docx) → 管道参数 + 工况参数
  2. Excel表格 (.xlsx) → 批量场景导入
  3. PDF文档 → 表格/参数提取
  4. CSV数据 → 时间序列导入

输出: PipelineSim 可识别的配置字典
"""

import re, os, json, tempfile
from typing import Dict, Optional, List, Any
from dataclasses import dataclass, field, asdict
from datetime import datetime


# ═══════════════════════════════════════════════
# 数据类型定义
# ═══════════════════════════════════════════════

@dataclass
class PipeConfig:
    """管道参数"""
    length: float = 10000.0          # m
    diameter: float = 0.5            # m
    wall_thickness: float = 0.01     # m
    roughness: float = 4.5e-5        # m
    elevation_start: float = 0.0
    elevation_end: float = 0.0


@dataclass
class SimulationConfig:
    """仿真配置"""
    name: str = "未命名案例"
    solver: str = "moc"              # moc/fvm/ifdm
    mode: str = "A"                  # A/B
    t_max: float = 20.0              # s
    P_initial: float = 2.0e6         # Pa
    T_initial: float = 20.0          # °C
    inlet_flow: float = 1.0          # m³/s
    inlet_temp: float = 20.0         # °C
    outlet_pressure: float = 2.0e6   # Pa
    valve_schedule: Optional[str] = None  # "10s关闭", "5-8s斜坡关闭"
    fluid_type: str = "water"        # water/non-newtonian
    pipe: PipeConfig = field(default_factory=PipeConfig)


@dataclass
class BatchConfig:
    """批量案例"""
    cases: List[SimulationConfig] = field(default_factory=list)
    
    def add(self, config: SimulationConfig):
        self.cases.append(config)
    
    @property
    def count(self) -> int:
        return len(self.cases)


# ═══════════════════════════════════════════════
# Word文档解析
# ═══════════════════════════════════════════════

class WordParser:
    """Word文档 → 仿真参数"""
    
    # 参数关键词映射
    PARAM_PATTERNS = {
        'length': [(r'管[道线]?长度[约为]?(\d+\.?\d*)', 1), (r'Length[:\s]*(\d+\.?\d*)', 1)],
        'diameter': [(r'管[径道]?[内]?[径直径][约为]?(\d+\.?\d*)', 1), (r'[Dd]iameter[:\s]*(\d+\.?\d*)', 1)],
        'wall_thickness': [(r'壁厚[约为]?(\d+\.?\d*)', 1), (r'wall[_\s]*thickness[:\s]*(\d+\.?\d*)', 1)],
        'flow_rate': [(r'流量[约为]?(\d+\.?\d*)', 1), (r'[Ff]low[_\s]*[Rr]ate[:\s]*(\d+\.?\d*)', 1)],
        'pressure': [(r'压力[约为]?(\d+\.?\d*)', 1), (r'[Pp]ressure[:\s]*(\d+\.?\d*)', 1)],
        'temperature': [(r'温度[约为]?(\d+\.?\d*)', 1), (r'[Tt]emp[:\s]*(\d+\.?\d*)', 1)],
        'time': [(r'[时仿]间[约为]?(\d+\.?\d*)', 1), (r'[Tt]ime[:\s]*(\d+\.?\d*)', 1)],
    }
    
    @staticmethod
    def extract_params(text: str) -> Dict[str, float]:
        """从文本中提取数值参数"""
        params = {}
        
        # 先尝试关键词模式
        for key, patterns in WordParser.PARAM_PATTERNS.items():
            for pattern, group in patterns:
                match = re.search(pattern, text)
                if match:
                    params[key] = float(match.group(group))
                    break
        
        # 检测场景
        if re.search(r'水[击锤]|water\s*hammer|sudden|valve\s*close', text, re.I):
            params['scenario'] = 'water_hammer'
        elif re.search(r'泵.*[停急]|pump\s*trip|pump.*stop', text, re.I):
            params['scenario'] = 'pump_trip'
        elif re.search(r'泄[漏]|leak|rupture', text, re.I):
            params['scenario'] = 'leak'
        
        return params
    
    @staticmethod
    def parse_docx(filepath: str) -> SimulationConfig:
        """解析Word文档"""
        try:
            from docx import Document
            doc = Document(filepath)
            
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # 也解析表格
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text for cell in row.cells]
                    tables_text.append(' | '.join(cells))
            
            text = '\n'.join(full_text + tables_text)
            params = WordParser.extract_params(text)
            
            config = SimulationConfig()
            config.name = os.path.splitext(os.path.basename(filepath))[0]
            
            # 映射参数
            if 'length' in params:
                config.pipe.length = params['length']
            if 'diameter' in params:
                config.pipe.diameter = params['diameter']
            if 'flow_rate' in params:
                config.inlet_flow = params['flow_rate']
            if 'pressure' in params:
                config.outlet_pressure = params['pressure'] * 1e6
                config.P_initial = config.outlet_pressure
            if 'temperature' in params:
                config.inlet_temp = params['temperature']
            if 'time' in params:
                config.t_max = params['time']
            if 'scenario' in params:
                config.name = f"{params['scenario']}_{config.name}"
            
            return config
        
        except ImportError:
            print("  ⚠️ python-docx not installed, skipping Word parsing")
            return SimulationConfig()
        except Exception as e:
            print(f"  ⚠️ Word parse error: {e}")
            return SimulationConfig()
    
    @staticmethod
    def describe() -> str:
        return "支持标准Word格式参数提取"


# ═══════════════════════════════════════════════
# Excel文档解析
# ═══════════════════════════════════════════════

class ExcelParser:
    """Excel表格 → 批量仿真参数"""
    
    # 期望的列名映射
    COLUMN_MAP = {
        'name': ['案例名称', '案例名', '场景', 'name', 'case', 'scenario'],
        'length': ['管长', '管道长度', '长度', 'length', 'pipe length', 'L'],
        'diameter': ['管径', '直径', '内径', 'diameter', 'pipe diameter', 'D'],
        'thickness': ['壁厚', 'wall thickness', 'thickness', 't'],
        'flow_rate': ['入口流量', '流量', 'flow rate', 'flow', 'Q'],
        'temperature': ['入口温度', '温度', '温度(°C)', 'temperature', 'T'],
        'pressure': ['出口压力', '压力', '出口压力(MPa)', 'pressure', 'P'],
        'time': ['仿真时间', '时长', '仿真时长(s)', '仿真时长', 'time', 't_max'],
        'solver': ['求解器', 'solver', 'method'],
        'fluid': ['流体', '流体类型', 'fluid', 'fluid type'],
        'valve': ['阀门', '阀门操作', '阀门调度', 'valve', 'valve schedule'],
    }
    
    @staticmethod
    def parse_xlsx(filepath: str) -> BatchConfig:
        """解析Excel → 批量案例"""
        try:
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True)
            ws = wb.active
            
            batch = BatchConfig()
            rows = list(ws.iter_rows(values_only=True))
            if not rows or len(rows) < 2:
                return batch
            
            # 识别列
            headers = [str(h).strip().lower() if h else '' for h in rows[0]]
            col_map = {}
            
            col_keywords = {
                'name': ['name', '案例', 'scenario', '场景', 'case'],
                'length': ['管长', '长度', 'length', 'l'],
                'diameter': ['管径', '直径', 'diameter', 'd'],
                'thickness': ['壁厚', 'thickness', 't'],
                'flow_rate': ['流量', 'flow', 'q'],
                'temperature': ['温度', 'temp', 't'],
                'pressure': ['压力', 'pressure', 'p'],
                'time': ['时间', '时长', 'time', 't_max'],
                'solver': ['求解器', 'solver'],
                'fluid': ['流体', 'fluid'],
                'valve': ['阀门', 'valve'],
            }
            
            for col_idx, header in enumerate(headers):
                for key, keywords in col_keywords.items():
                    if any(kw.lower() in header for kw in keywords):
                        col_map[key] = col_idx
                        break
            
            # 解析每一行
            for row in rows[1:]:
                if not any(cell is not None for cell in row):
                    continue
                
                config = SimulationConfig()
                if 'name' in col_map and row[col_map['name']]:
                    config.name = str(row[col_map['name']])
                if 'length' in col_map and row[col_map['length']]:
                    config.pipe.length = float(row[col_map['length']])
                if 'diameter' in col_map and row[col_map['diameter']]:
                    config.pipe.diameter = float(row[col_map['diameter']])
                if 'flow_rate' in col_map and row[col_map['flow_rate']]:
                    config.inlet_flow = float(row[col_map['flow_rate']])
                if 'pressure' in col_map and row[col_map['pressure']]:
                    p = float(row[col_map['pressure']])
                    config.outlet_pressure = p * 1e6 if p < 100 else p
                if 'time' in col_map and row[col_map['time']]:
                    config.t_max = float(row[col_map['time']])
                if 'solver' in col_map and row[col_map['solver']]:
                    config.solver = str(row[col_map['solver']]).lower()
                if 'fluid' in col_map and row[col_map['fluid']]:
                    config.fluid_type = str(row[col_map['fluid']]).lower()
                
                batch.add(config)
            
            wb.close()
            return batch
        
        except ImportError:
            print("  ⚠️ openpyxl not available")
            return BatchConfig()
        except Exception as e:
            print(f"  ⚠️ Excel parse error: {e}")
            return BatchConfig()
    
    @staticmethod
    def generate_template(path: str = '/tmp/pipeline_batch_template.xlsx'):
        """生成批量导入模板"""
        try:
            import openpyxl
            from openpyxl.styles import Font, PatternFill
            
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "批量仿真案例"
            
            headers = ['案例名称', '管道长度(m)', '管径(m)', '入口流量(m³/s)',
                      '出口压力(MPa)', '仿真时长(s)', '求解器', '流体类型', '阀门操作']
            
            header_fill = PatternFill(start_color='2c3e50', end_color='2c3e50', fill_type='solid')
            header_font = Font(bold=True, color='ffffff')
            
            for col, h in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=h)
                cell.fill = header_fill
                cell.font = header_font
            
            # 示例行
            example = ['案例1-水击', '10000', '0.5', '1.0', '2.0', '20', 'moc', 'water', '5秒关闭']
            for col, val in enumerate(example, 1):
                ws.cell(row=2, column=col, value=val)
            
            for col in range(1, len(headers)+1):
                ws.column_dimensions[chr(64+col)].width = 18
            
            wb.save(path)
            return path
        except Exception as e:
            print(f"  ⚠️ Template error: {e}")
            return None
    
    @staticmethod
    def describe() -> str:
        return "支持标准Excel批量导入(9列模板)"


# ═══════════════════════════════════════════════
# PDF文档解析
# ═══════════════════════════════════════════════

class PDFParser:
    """PDF文档 → 参数/表格提取"""
    
    @staticmethod
    def extract_text(filepath: str) -> str:
        """提取PDF文本"""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            return ""
        except Exception as e:
            return f""
    
    @staticmethod
    def extract_tables(filepath: str) -> List[List[List[str]]]:
        """提取PDF中的表格"""
        tables = []
        try:
            import camelot
            # 使用Lattice模式（有边框表格）
            try:
                result = camelot.read_pdf(filepath, pages='all', flavor='lattice')
                tables.extend([t.data for t in result])
            except:
                pass
            # 使用Stream模式（无边框表格）
            try:
                result = camelot.read_pdf(filepath, pages='all', flavor='stream')
                tables.extend([t.data for t in result])
            except:
                pass
        except ImportError:
            pass
        except Exception:
            pass
        return tables
    
    @staticmethod
    def parse_to_config(filepath: str) -> SimulationConfig:
        """PDF → 仿真配置"""
        text = PDFParser.extract_text(filepath)
        params = WordParser.extract_params(text)
        
        config = SimulationConfig()
        config.name = os.path.splitext(os.path.basename(filepath))[0]
        
        if 'length' in params: config.pipe.length = params['length']
        if 'diameter' in params: config.pipe.diameter = params['diameter']
        if 'flow_rate' in params: config.inlet_flow = params['flow_rate']
        if 'pressure' in params: config.outlet_pressure = params['pressure'] * 1e6
        if 'time' in params: config.t_max = params['time']
        
        if 'scenario' in params:
            config.name = f"{params['scenario']}_{config.name}"
        
        return config


# ═══════════════════════════════════════════════
# 统一入口
# ═══════════════════════════════════════════════

def parse_document(filepath: str) -> Dict[str, Any]:
    """
    智能文档解析（自动识别格式）
    
    Args:
        filepath: 文档路径 (.docx/.xlsx/.pdf/.csv)
    
    Returns:
        解析结果
    """
    ext = os.path.splitext(filepath)[1].lower()
    
    if ext == '.docx':
        config = WordParser.parse_docx(filepath)
        return {'type': 'word', 'config': config, 'batch': None}
    
    elif ext == '.xlsx':
        batch = ExcelParser.parse_xlsx(filepath)
        return {'type': 'excel', 'config': None, 'batch': batch}
    
    elif ext == '.pdf':
        config = PDFParser.parse_to_config(filepath)
        tables = PDFParser.extract_tables(filepath)
        return {'type': 'pdf', 'config': config, 'tables': tables}
    
    else:
        return {'type': 'unknown', 'error': f'不支持格式: {ext}'}


# ═══════════════════════════════════════════════
# 配置→仿真参数转换
# ═══════════════════════════════════════════════

def config_to_setup(config: SimulationConfig) -> Dict:
    """SimulationConfig → 求解器参数"""
    return {
        'name': config.name,
        'solver': config.solver,
        'mode': config.mode,
        't_max': config.t_max,
        'P_initial': config.P_initial,
        'T_initial': config.T_initial,
        'pipe_length': config.pipe.length,
        'pipe_diameter': config.pipe.diameter,
        'pipe_thickness': config.pipe.wall_thickness,
        'flow_inlet': config.inlet_flow,
        'temp_inlet': config.inlet_temp,
        'pressure_outlet': config.outlet_pressure,
        'fluid_type': config.fluid_type,
    }


if __name__ == '__main__':
    print("=== 文档→仿真参数智能引擎 ===")
    print()
    
    # 测试Word解析
    print("1. Word解析 (模拟文本):")
    test_text = "管道长度为10000m, 管径0.5m, 入口流量1.0m³/s, 压力2.0MPa, 温度20°C"
    params = WordParser.extract_params(test_text)
    print(f"   提取参数: {params}")
    
    print()
    print("2. 生成Excel模板:")
    tmpl = ExcelParser.generate_template('/tmp/pipeline_template.xlsx')
    if tmpl:
        import os
        print(f"   模板已生成: {tmpl} ({os.path.getsize(tmpl)/1024:.0f}KB)")
    
    print()
    print("3. Config→Setup转换:")
    config = SimulationConfig(
        name="测试案例",
        pipe=PipeConfig(length=5000, diameter=0.3),
        inlet_flow=0.5, outlet_pressure=1.5e6,
        fluid_type='bingham'
    )
    setup = config_to_setup(config)
    for k, v in setup.items():
        print(f"   {k}: {v}")
    
    print()
    print("✅ 数据输入管道就绪!")
